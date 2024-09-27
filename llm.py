# import streamlit as st
# import fitz  # PyMuPDF
# import docx
# import openai
# import re

# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# def extract_text_from_word(file):
#     doc = docx.Document(file)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text

# def extract_information(text):
#     openai.api_key = 'your_openai_api_key'
#     prompt = f"Extract the following details from the text:\n\n{text}\n\nPersonal Information:\n- Name\n- Gender\n- Age\nEducation:\n- Academic level\n- Institution\n- GPA/Grade\n- Start date – End date\nWork Experience:\n- Company\n- Location\n- Role\n- Start date – End date\n- Description"
    
#     response = openai.Completion.create(
#         engine="text-davinci-003",
#         prompt=prompt,
#         max_tokens=500
#     )
    
#     return response.choices.text.strip()

# st.title("Document Information Extractor")

# uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

# if uploaded_file is not None:
#     if uploaded_file.type == "application/pdf":
#         text = extract_text_from_pdf(uploaded_file)
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         text = extract_text_from_word(uploaded_file)
    
#     st.write("Extracted Text:")
#     st.write(text)
    
#     if st.button("Extract Information"):
#         extracted_info = extract_information(text)
#         st.write("Extracted Information:")
#         st.write(extracted_info)
export KMP_DUPLICATE_LIB_OK=TRUE

import streamlit as st
import fitz  # PyMuPDF
import docx
from transformers import pipeline

def extract_text_from_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_word(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_information(text):
    nlp = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english")
    ner_results = nlp(text)
    
    extracted_info = {
        "Personal Information": {},
        "Education": [],
        "Work Experience": []
    }
    
    # Process NER results to extract relevant information
    for entity in ner_results:
        if entity['entity'] == 'B-PER':
            extracted_info["Personal Information"]["Name"] = entity['word']
        # Add more conditions to extract other details
    
    return extracted_info

st.title("Document Information Extractor")

uploaded_file = st.file_uploader("Upload a PDF or Word document", type=["pdf", "docx"])

if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_word(uploaded_file)
    
    st.write("Extracted Text:")
    st.write(text)
    
    if st.button("Extract Information"):
        extracted_info = extract_information(text)
        st.write("Extracted Information:")
        st.write(extracted_info)

