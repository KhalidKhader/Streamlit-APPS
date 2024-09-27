import streamlit as st
import openai
import PyPDF2
import docx
import os


# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                full_text.append(para_text)

        # Extract text from tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return None

# Function to call GPT-4-turbo for information extraction
def extract_information_with_gpt(text):
    prompt = f"""
    Extract the following details from this CV:
    - Name
    - Gender
    - Age
    - Education (Degree, Institution, GPA, Start Date, End Date)
    - Work Experience (Company, Role, Start Date, End Date, Description)

    CV Text:
    {text}
    """
    
    # Call GPT-4-turbo
    response = openai.ChatCompletion.create(
         model= "gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that extracts information from CVs."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=1000,
        temperature=0.2
    )
    
    return response['choices'][0]['message']['content'].strip()

# Streamlit UI
st.title("CV Information Extractor")

uploaded_file = st.file_uploader("Upload your CV (PDF or Word)", type=["pdf", "docx"])

if uploaded_file is not None:
    file_extension = uploaded_file.name.split(".")[-1].lower()

    # Extract text based on file type
    if file_extension == "pdf":
        text = extract_text_from_pdf(uploaded_file)
    elif file_extension == "docx":
        text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a PDF or Word document.")
        text = None

    if text:
        st.subheader("Extracted CV Text")
        st.write(text)

        # Extract information using GPT-4-turbo
        with st.spinner("Extracting information with GPT-4-turbo..."):
            extracted_info = extract_information_with_gpt(text)
            st.subheader("Extracted Information")
            st.write(extracted_info)
