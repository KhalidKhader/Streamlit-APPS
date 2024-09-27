import streamlit as st
import openai
import PyPDF2
import docx
import os

openai.api_key = os.getenv('openai_api_key')
# Set your OpenAI API key here
openai.api_key = "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" # TO-DO

# # Function to extract text from PDF
# def extract_text_from_pdf(pdf_file):
#     pdf_reader = PyPDF2.PdfReader(pdf_file)
#     text = ''
#     for page_num in range(len(pdf_reader.pages)):
#         page = pdf_reader.pages[page_num]
#         text += page.extract_text()
#     return text


# def extract_text_from_docx(docx_file):
#     try:
#         doc = docx.Document(docx_file)
#         full_text = []

#         # Extract text from paragraphs
#         for para in doc.paragraphs:
#             para_text = para.text.strip()
#             if para_text:
#                 full_text.append(para_text)

#         # Extract text from tables if they exist
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     cell_text = cell.text.strip()
#                     if cell_text:
#                         full_text.append(cell_text)

#         return '\n'.join(full_text)
#     except Exception as e:
#         print(f"Error reading DOCX file: {e}")
#         return None



# # Function to call GPT-4-turbo for information extraction
# def extract_information_with_gpt(text):
#     prompt = f"""
#     Extract the following details from this CV:
#     - Name
#     - Gender
#     - Age
#     - Education (Degree, Institution, GPA, Start Date, End Date)
#     - Work Experience (Company, Role, Start Date, End Date, Description)

#     CV Text:
#     {text}
#     """
    
#     # Call GPT-4-turbo
#     response = openai.ChatCompletion.create(
#          model= "gpt-4o-mini",
#         messages=[
#             {"role": "system", "content": "You are a helpful assistant that extracts information from CVs."},
#             {"role": "user", "content": prompt},
#         ],
#         max_tokens=1000,
#         temperature=0.2
#     )
    
#     return response['choices'][0]['message']['content'].strip()

# # Streamlit UI
# st.title("CV Information Extractor")

# uploaded_file = st.file_uploader("Upload your CV (PDF or Word)", type=["pdf", "docx"])

# if uploaded_file is not None:
#     file_extension = uploaded_file.name.split(".")[-1].lower()

#     # Extract text based on file type
#     if file_extension == "pdf":
#         text = extract_text_from_pdf(uploaded_file)
#     elif file_extension == "docx":
#         text = extract_text_from_docx(uploaded_file)
#     else:
#         st.error("Unsupported file type. Please upload a PDF or Word document.")
#         text = None

#     if text:
#         st.subheader("Extracted CV Text")
#         st.write(text)

#         # Extract information using GPT-4-turbo
#         with st.spinner("Extracting information with GPT-4-turbo..."):
#             extracted_info = extract_information_with_gpt(text)
#             st.subheader("Extracted Information")
#             st.write(extracted_info)



# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        return '\n'.join(_extract_paragraphs_and_tables(doc))
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

# Helper function to extract text from DOCX paragraphs and tables
def _extract_paragraphs_and_tables(doc):
    full_text = []

    # Extract paragraphs
    _extract_paragraphs(doc, full_text)
    
    # Extract tables
    _extract_tables(doc, full_text)
    
    return full_text

# Extract text from paragraphs in DOCX
def _extract_paragraphs(doc, full_text):
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            full_text.append(para_text)

# Extract text from tables in DOCX
def _extract_tables(doc, full_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)

# Function to select the correct extractor based on file extension
def extract_text_based_on_extension(file, file_extension):
    if file_extension == "pdf":
        return extract_text_from_pdf(file)
    elif file_extension == "docx":
        return extract_text_from_docx(file)
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")
        return None

# Function to call GPT-4 for information extraction
def extract_information_with_gpt(text):
    try:
        prompt = _create_gpt_prompt(text)
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that extracts information from CVs."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1000,
            temperature=0.2
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        st.error(f"Error with GPT extraction: {e}")
        return None

# Helper function to create GPT prompt
def _create_gpt_prompt(text):
    return f"""
    Extract the following details from this CV:
    - Name
    - Gender
    - Age 
    - Education (Degree, Institution, GPA, Start Date, End Date)
    - Work Experience (Company, Role, Start Date, End Date, Description)

    CV Text:
    {text}
    """

# Main function to handle the Streamlit UI
def main():
    st.title("CV Information Extractor")

    uploaded_file = st.file_uploader("Upload your CV (PDF or Word)", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()

        # Extract text based on file type
        text = extract_text_based_on_extension(uploaded_file, file_extension)

        if text:
            st.subheader("Extracted CV Text")
            st.write(text)

            # Extract information using GPT
            with st.spinner("Extracting information with GPT..."):
                extracted_info = extract_information_with_gpt(text)
                if extracted_info:
                    st.subheader("Extracted Information")
                    st.write(extracted_info)


if __name__ == "__main__":
    main()
