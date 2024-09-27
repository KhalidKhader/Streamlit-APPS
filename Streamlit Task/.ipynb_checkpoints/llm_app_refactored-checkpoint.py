import streamlit as st
import openai
import PyPDF2
import docx
import os

openai.api_key = os.getenv('openai_api_key')

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        return '\n'.join(_extract_paragraphs_and_tables(doc))
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

# Helper function to extract text from DOCX paragraphs and tables
def _extract_paragraphs_and_tables(doc):
    full_text = []
    # Extract paragraphs
    _extract_paragraphs(doc, full_text)
    # Extract tables
    _extract_tables(doc, full_text)
    return full_text

# Extract text from paragraphs in DOCX
def _extract_paragraphs(doc, full_text):
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            full_text.append(para_text)

# Extract text from tables in DOCX
def _extract_tables(doc, full_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)

# Function to select the correct extractor based on file extension
def extract_text_based_on_extension(file, file_extension):
    if file_extension == "pdf":
        return extract_text_from_pdf(file)
    elif file_extension == "docx":
        return extract_text_from_docx(file)
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")
        return None

# Function to call GPT-4 for information extraction
def extract_information_with_gpt(text):
    try:
        prompt = _create_gpt_prompt(text)
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that extracts information from CVs."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1000,
            temperature=0.2
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        st.error(f"Error with GPT extraction: {e}")
        return None

# Helper function to create GPT prompt
def _create_gpt_prompt(text):
    return f"""
    Extract the following details from this CV:
    - Name
    - Gender
    - Age 
    - Education (Degree, Institution, GPA, Start Date, End Date)
    - Work Experience (Company, Role, Start Date, End Date, Description)

    CV Text:
    {text}
    """

# Main function to handle the Streamlit UI
def main():
    st.title("CV Information Extractor")

    uploaded_file = st.file_uploader("Upload your CV (PDF or Word)", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()

        # Extract text based on file type
        text = extract_text_based_on_extension(uploaded_file, file_extension)

        if text:
            st.subheader("Extracted CV Text")
            st.write(text)

            # Extract information using GPT
            with st.spinner("Extracting information with GPT..."):
                extracted_info = extract_information_with_gpt(text)
                if extracted_info:
                    st.subheader("Extracted Information")
                    st.write(extracted_info)


if __name__ == "__main__":
    main()
